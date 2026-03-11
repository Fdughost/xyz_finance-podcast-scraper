#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
播客监控日报生成模块
生成Word格式的每日播客数据报告
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from typing import List, Dict
import logging

from data_manager import PodcastDataManager

logger = logging.getLogger(__name__)


class ReportGenerator:
    """日报生成器类"""
    
    def __init__(self, data_manager: PodcastDataManager):
        """
        初始化日报生成器
        
        Args:
            data_manager: 数据管理器实例
        """
        self.data_manager = data_manager
    
    def generate_daily_report(self, output_file: str = None) -> str:
        """
        生成每日播客监控报告
        
        Args:
            output_file: 输出文件路径，默认为当前日期
            
        Returns:
            生成的报告文件路径
        """
        if output_file is None:
            today = datetime.now().strftime('%Y-%m-%d')
            output_file = f'reports/播客监控日报_{today}.docx'
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 设置中文字体
            doc.styles['Normal'].font.name = '微软雅黑'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            
            # 添加标题
            self._add_title(doc)
            
            # 添加报告日期
            self._add_report_info(doc)
            
            # 获取对比数据
            comparison_report = self.data_manager.get_comparison_report()
            
            if not comparison_report:
                self._add_no_data_message(doc)
            else:
                # 添加概览
                self._add_overview(doc, comparison_report)
                
                # 添加详细数据表格
                self._add_detailed_table(doc, comparison_report)
                
                # 添加分类统计
                self._add_category_summary(doc, comparison_report)
            
            # 添加页脚
            self._add_footer(doc)
            
            # 保存文档
            doc.save(output_file)
            logger.info(f"日报生成成功: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"生成日报失败: {e}")
            raise
    
    def _add_title(self, doc: Document):
        """添加报告标题"""
        from datetime import datetime
        today = datetime.now().strftime('%Y-%m-%d')
        title = doc.add_heading(f'播客监控日报 {today}', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置标题格式
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    def _add_report_info(self, doc: Document):
        """添加报告信息"""
        today = datetime.now()
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_run = info_para.add_run(f'报告日期: {today.strftime("%Y年%m月%d日")}')
        info_run.font.size = Pt(12)
        info_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # 空行
    
    def _add_overview(self, doc: Document, comparison_report: List[Dict]):
        """添加数据概览"""
        doc.add_heading('一、数据概览', level=1)
        
        # 统计数据
        total_podcasts = len(comparison_report)
        new_episodes = sum(1 for item in comparison_report if item['changes'].get('has_new_episode', False))
        total_subscriber_change = sum(item['changes'].get('subscribers', 0) for item in comparison_report)
        
        # 添加概览段落
        overview_items = [
            f"📊 监控播客总数: {total_podcasts} 个",
            f"🆕 今日更新节目: {new_episodes} 个",
            f"👥 总订阅数变化: {total_subscriber_change:+,d} 人"
        ]
        
        for item in overview_items:
            para = doc.add_paragraph(item, style='List Bullet')
            para.runs[0].font.size = Pt(11)
        
        doc.add_paragraph()  # 空行
    
    def _add_detailed_table(self, doc: Document, comparison_report: List[Dict]):
        """添加详细数据表格"""
        doc.add_heading('二、详细数据', level=1)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # 设置表头
        header_cells = table.rows[0].cells
        headers = ['播客名称', '分类', '当前订阅数', '订阅数变化', '是否更新', '最新节目']
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            
            # 设置表头格式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 设置表头背景色
            cell._element.get_or_add_tcPr().append(
                self._create_cell_shading((0, 51, 102))
            )
        
        # 添加数据行
        for item in comparison_report:
            row_cells = table.add_row().cells
            
            # 播客名称
            row_cells[0].text = item.get('podcast_name', '未知')
            
            # 分类
            latest_snapshot = self.data_manager.get_latest_snapshot(item['podcast_id'])
            category = latest_snapshot.get('category', '未分类') if latest_snapshot else '未分类'
            row_cells[1].text = category
            
            # 当前订阅数
            current_subscribers = item['today'].get('subscribers', 0)
            row_cells[2].text = f"{current_subscribers:,d}"
            
            # 订阅数变化
            subscriber_change = item['changes'].get('subscribers', 0)
            change_text = f"{subscriber_change:+,d}"
            row_cells[3].text = change_text
            
            # 设置变化数字的颜色
            if subscriber_change > 0:
                for paragraph in row_cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
            elif subscriber_change < 0:
                for paragraph in row_cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            
            # 是否更新
            has_new = item['changes'].get('has_new_episode', False)
            row_cells[4].text = '✓ 是' if has_new else '✗ 否'
            
            # 最新节目标题
            if latest_snapshot:
                latest_title = latest_snapshot.get('latest_episode_title', '')
                # 截取标题前50个字符
                if len(latest_title) > 50:
                    latest_title = latest_title[:50] + '...'
                row_cells[5].text = latest_title
            else:
                row_cells[5].text = '-'
            
            # 设置单元格格式
            for i, cell in enumerate(row_cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    if i in [2, 3, 4]:  # 居中对齐
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
    
    def _add_category_summary(self, doc: Document, comparison_report: List[Dict]):
        """添加分类统计"""
        doc.add_heading('三、分类统计', level=1)
        
        # 按分类统计
        category_stats = {}
        for item in comparison_report:
            latest_snapshot = self.data_manager.get_latest_snapshot(item['podcast_id'])
            category = latest_snapshot.get('category', '未分类') if latest_snapshot else '未分类'
            
            if category not in category_stats:
                category_stats[category] = {
                    'count': 0,
                    'total_subscribers': 0,
                    'new_episodes': 0
                }
            
            category_stats[category]['count'] += 1
            category_stats[category]['total_subscribers'] += item['today'].get('subscribers', 0)
            if item['changes'].get('has_new_episode', False):
                category_stats[category]['new_episodes'] += 1
        
        # 创建统计表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light List Accent 1'
        
        # 表头
        header_cells = table.rows[0].cells
        headers = ['分类', '播客数量', '总订阅数', '今日更新数']
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 数据行
        for category, stats in sorted(category_stats.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = category
            row_cells[1].text = str(stats['count'])
            row_cells[2].text = f"{stats['total_subscribers']:,d}"
            row_cells[3].text = str(stats['new_episodes'])
            
            for i, cell in enumerate(row_cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    if i > 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
    
    def _add_no_data_message(self, doc: Document):
        """添加无数据提示"""
        para = doc.add_paragraph('暂无监控数据，请先运行爬虫程序采集数据。')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(12)
        para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    def _add_footer(self, doc: Document):
        """添加页脚"""
        doc.add_page_break()
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_text = f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n'
        footer_text += '本报告由播客监控系统自动生成'
        
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    def _create_cell_shading(self, rgb_tuple):
        """创建单元格底色"""
        from docx.oxml import OxmlElement
        
        shading_elm = OxmlElement('w:shd')
        r, g, b = rgb_tuple
        shading_elm.set(qn('w:fill'), f'{r:02X}{g:02X}{b:02X}')
        return shading_elm


def main():
    """主函数"""
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    logger.info("开始生成播客监控日报")
    
    # 创建数据管理器
    data_manager = PodcastDataManager()
    
    # 创建报告生成器
    report_generator = ReportGenerator(data_manager)
    
    # 生成日报
    report_file = report_generator.generate_daily_report()
    
    logger.info(f"日报生成完成: {report_file}")


if __name__ == '__main__':
    main()
